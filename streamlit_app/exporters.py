import os
import sys
from pathlib import Path
import csv
import io

# Setup python path to include the current directory so we can import backend
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# Try to import backend services
try:
    from backend.app.services.generation.pptx_builder import build_pptx
    from backend.app.services.generation.pdf_builder import build_pdf
    from backend.app.schemas.slide import SlideContent
except ImportError:
    build_pptx = None
    build_pdf = None
    SlideContent = None

# ── Word (DOCX) Exporter ───────────────────────────────────────────────────
def export_to_docx(content_items, title="Content Curator Deliverable"):
    """
    Generates a DOCX file using python-docx and returns bytes.
    content_items can be:
      - a list of dicts with 'title' and 'content' / 'bullets'
      - a list of slides
      - raw text
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # Return none or raise
        return None

    doc = Document()
    
    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Title (Tata Steel branded styling)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1E, 0x27, 0x61) # Tata Blue
    title_p.paragraph_format.space_after = Pt(24)

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Tata Steel Content Curator · Automated Deliverable")
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    sub_p.paragraph_format.space_after = Pt(36)

    # Add content
    if isinstance(content_items, str):
        doc.add_paragraph(content_items)
    else:
        for idx, item in enumerate(content_items):
            # Check structure (mock lists or slide structures)
            heading_text = item.get('title', item.get('heading', f"Section {idx+1}"))
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
            h_run = h.add_run(heading_text)
            h_run.font.color.rgb = RGBColor(0x1E, 0x27, 0x61) # Tata Blue

            # Bullets or content
            if 'bullets' in item:
                for bullet in item['bullets']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.space_after = Pt(4)
                    p.add_run(bullet)
            elif 'content' in item:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.15
                p.add_run(item['content'])
            elif 'subtitle' in item: # Title slide handling
                p = doc.add_paragraph()
                p.add_run(item['subtitle']).font.italic = True

    # Save to memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

# ── Excel (XLSX) Exporter ───────────────────────────────────────────────────
def export_to_xlsx(tabular_data, sheet_name="Deliverable"):
    """
    Generates an XLSX spreadsheet bytes from list of dicts.
    If openpyxl is missing, falls back to generating a CSV.
    """
    if not tabular_data:
        return b""

    # Try using openpyxl for formatted Excel sheets
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name[:30] # Excel limit is 31 chars
        
        # Gridlines visible
        ws.views.sheetView[0].showGridLines = True
        
        # Determine headers
        if hasattr(tabular_data[0], 'keys'):
            headers = list(tabular_data[0].keys())
        else:
            headers = [f"Column {i}" for i in range(len(tabular_data[0]))]

        # Write header row
        ws.append([str(h).upper() for h in headers])
        
        # Style headers
        tata_blue_fill = PatternFill(start_color="1E2761", end_color="1E2761", fill_type="solid")
        header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
        center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        thin_side = Side(border_style="thin", color="CCCCCC")
        border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)

        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.fill = tata_blue_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = border
            
        ws.row_dimensions[1].height = 28

        # Write data rows
        row_font = Font(name="Calibri", size=10)
        left_align = Alignment(horizontal="left", vertical="center")
        
        for row_idx, row_dict in enumerate(tabular_data, start=2):
            row_values = []
            for h in headers:
                val = row_dict.get(h, "") if hasattr(row_dict, 'get') else ""
                row_values.append(val)
            
            ws.append(row_values)
            
            # Style each row
            for col_idx in range(1, len(headers) + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.font = row_font
                cell.alignment = left_align
                cell.border = border
                
            ws.row_dimensions[row_idx].height = 20

        # Auto-fit columns
        for col in ws.columns:
            max_len = 0
            for cell in col:
                val_str = str(cell.value or '')
                if len(val_str) > max_len:
                    max_len = len(val_str)
            col_letter = openpyxl.utils.get_column_letter(col[0].column)
            ws.column_dimensions[col_letter].width = max(max_len + 3, 12)

        file_stream = io.BytesIO()
        wb.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()

    except ImportError:
        # Fallback to CSV format but serve it as excel-compatible CSV
        output = io.StringIO()
        if hasattr(tabular_data[0], 'keys'):
            headers = list(tabular_data[0].keys())
        else:
            headers = []
            
        writer = csv.writer(output)
        if headers:
            writer.writerow(headers)
        for row in tabular_data:
            if hasattr(row, 'get'):
                writer.writerow([row.get(h, "") for h in headers])
            else:
                writer.writerow(row)
                
        return output.getvalue().encode('utf-8')

# ── PPTX & PDF Wrappers ─────────────────────────────────────────────────────
def export_to_pptx(slides_list, title="Content Curator Deliverable", theme="midnight_executive"):
    """
    Generates a PPTX file using backend python-pptx services.
    Returns file bytes.
    """
    if build_pptx is None or SlideContent is None:
        # In case backend packages are not imported, return dummy or mock
        return None
        
    formatted_slides = []
    for s in slides_list:
        slide_obj = SlideContent(
            slide_number=s.get('id', s.get('slide_number', 1)),
            title=s.get('title', 'Untitled Slide'),
            bullets=s.get('bullets', []),
            layout=s.get('type', 'content'),
            speaker_notes=s.get('speaker_notes', '')
        )
        formatted_slides.append(slide_obj)
        
    pptx_path = build_pptx(formatted_slides, title, theme)
    if pptx_path.exists():
        with open(pptx_path, "rb") as f:
            data = f.read()
        try:
            os.remove(pptx_path) # Cleanup
        except OSError:
            pass
        return data
    return None

def export_to_pdf(slides_list, title="Content Curator Deliverable", theme="midnight_executive"):
    """
    Generates a PDF file using backend ReportLab services.
    Returns file bytes.
    """
    if build_pdf is None or SlideContent is None:
        return None
        
    formatted_slides = []
    for s in slides_list:
        slide_obj = SlideContent(
            slide_number=s.get('id', s.get('slide_number', 1)),
            title=s.get('title', 'Untitled Slide'),
            bullets=s.get('bullets', []),
            layout=s.get('type', 'content'),
            speaker_notes=s.get('speaker_notes', '')
        )
        formatted_slides.append(slide_obj)
        
    pdf_path = build_pdf(formatted_slides, title, theme)
    if pdf_path.exists():
        with open(pdf_path, "rb") as f:
            data = f.read()
        try:
            os.remove(pdf_path) # Cleanup
        except OSError:
            pass
        return data
    return None
