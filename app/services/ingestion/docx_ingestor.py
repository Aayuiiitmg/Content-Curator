from pathlib import Path
from docx import Document


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from a DOCX file, preserving paragraph structure."""
    doc = Document(str(file_path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)
